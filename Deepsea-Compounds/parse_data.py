#!/usr/bin/env python3
"""
Parse Deepsea-Compounds docx files and generate JSON datasets.
Run: python3 parse_data.py
"""

import json
import os
from docx import Document

# Paths — update these to point to your docx files
TABLE1_PATH = "/mnt/d/github_local/Deepsea-Compounds-main/Deepsea-Compounds-main/表1 深海深渊化合物的80种靶标信息.docx"
TABLE2_PATH = "/mnt/d/github_local/Deepsea-Compounds-main/Deepsea-Compounds-main/表2 324种深海来源化合物信息.docx"
OUTPUT_DIR = "/home/wcf/Deepsea-Compounds/assets/data"

# ─────────────────────────────────────────────────────────────────────────────
# TRY_PATHS: try multiple possible locations for the docx files
# ─────────────────────────────────────────────────────────────────────────────
import os, glob

def try_find(paths):
    for p in paths:
        if os.path.exists(p):
            return p
    return paths[0]

TABLE1_PATH = try_find([
    "/mnt/d/github_local/Deepsea-Compounds-main/Deepsea-Compounds-main/表1 深海深渊化合物的80种靶标信息.docx",
    os.path.expanduser("~/Deepsea-Compounds/表1 深海深渊化合物的80种靶标信息.docx"),
    "./表1 深海深渊化合物的80种靶标信息.docx",
])
TABLE2_PATH = try_find([
    "/mnt/d/github_local/Deepsea-Compounds-main/Deepsea-Compounds-main/表2 324种深海来源化合物信息.docx",
    os.path.expanduser("~/Deepsea-Compounds/表2 324种深海来源化合物信息.docx"),
    "./表2 324种深海来源化合物信息.docx",
])

# ─────────────────────────────────────────────────────────────────────────────
# 1. Parse Table 2: 324 Deep-sea Compounds
#    Layout: 6 columns × ~111 rows
#    Cols 0-1: ID 1-100 + Name
#    Cols 2-3: ID 101-200 + Name
#    Cols 4-5: ID 201-324 + Name
# ─────────────────────────────────────────────────────────────────────────────
def parse_compounds(doc_path):
    doc = Document(doc_path)
    table = doc.tables[0]

    compounds = {}
    for row in table.rows[1:]:   # skip header row
        cells = row.cells
        # Extract all three ID→Name pairs across the row
        for offset in range(0, 6, 2):
            id_text  = cells[offset].text.strip()
            name     = cells[offset + 1].text.strip()
            if id_text.isdigit() and name:
                compounds[int(id_text)] = name

    # Sort by ID and convert to list
    compounds_list = [{"id": k, "name": v} for k, v in sorted(compounds.items())]
    print(f"  Compounds parsed: {len(compounds_list)}")
    return compounds_list


# ─────────────────────────────────────────────────────────────────────────────
# 2. Parse Table 1: 80 Targets
#    Columns: 编号 | 靶标名称 | Uniport | Example of deep-sea compounds | Assays | Reference
# ─────────────────────────────────────────────────────────────────────────────
def parse_targets(doc_path):
    doc = Document(doc_path)
    table = doc.tables[0]

    targets = []
    for row in table.rows[1:]:   # skip header row
        cells = row.cells
        targets.append({
            "id":       cells[0].text.strip(),
            "name":     cells[1].text.strip(),
            "uniprot":  cells[2].text.strip(),
            "examples": cells[3].text.strip(),   # free-text compound hints
            "assays":   cells[4].text.strip(),
            "reference":cells[5].text.strip()
        })
    print(f"  Targets parsed: {len(targets)}")
    return targets


# ─────────────────────────────────────────────────────────────────────────────
# 3. Build associations: fuzzy-match target.example_compounds against compound names
#    Returns a list of {target_id, compound_ids, unmatched_examples}
# ─────────────────────────────────────────────────────────────────────────────
def fuzzy_match(query, compounds, threshold=0.6):
    """Return list of (compound_id, score) sorted by descending score."""
    results = []
    q = query.lower()
    for c in compounds:
        name = c["name"].lower()
        # Exact substring match → high score
        if q in name or name in q:
            score = 1.0
        else:
            # Simple token-level overlap
            q_tokens = set(q.replace("-", " ").split())
            n_tokens = set(name.replace("-", " ").split())
            if q_tokens and n_tokens:
                overlap = len(q_tokens & n_tokens) / max(len(q_tokens), len(n_tokens))
                score = overlap
            else:
                score = 0
        if score >= threshold:
            results.append((c["id"], score))
    return sorted(results, key=lambda x: -x[1])


def build_associations(targets, compounds):
    associations = []
    for t in targets:
        matched = fuzzy_match(t["examples"], compounds, threshold=0.5)
        compound_ids = [cid for cid, score in matched]
        associations.append({
            "target_id":          int(t["id"]),
            "target_name":        t["name"],
            "matched_compound_ids": compound_ids,
            "example_text":       t["examples"]
        })
    return associations


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────
def main():
    print("Parsing Table 2 (Compounds)...")
    compounds = parse_compounds(TABLE2_PATH)

    print("Parsing Table 1 (Targets)...")
    targets = parse_targets(TABLE1_PATH)

    print("Building associations...")
    associations = build_associations(targets, compounds)

    # Backfill compound → target links
    compound_target_map = {c["id"]: [] for c in compounds}
    for a in associations:
        for cid in a["matched_compound_ids"]:
            if cid in compound_target_map:
                compound_target_map[cid].append(a["target_id"])

    compounds_out = []
    for c in compounds:
        compounds_out.append({
            "id":             c["id"],
            "name":           c["name"],
            "target_ids":     compound_target_map.get(c["id"], [])
        })

    targets_out = []
    for t in targets:
        matched_ids = [a["matched_compound_ids"] for a in associations
                       if int(a["target_id"]) == int(t["id"])]
        targets_out.append({
            "id":              int(t["id"]),
            "name":            t["name"],
            "uniprot":         t["uniprot"],
            "assays":          t["assays"],
            "reference":       t["reference"],
            "compound_ids":    matched_ids[0] if matched_ids else []
        })

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    with open(f"{OUTPUT_DIR}/compounds.json", "w", encoding="utf-8") as f:
        json.dump(compounds_out, f, ensure_ascii=False, indent=2)

    with open(f"{OUTPUT_DIR}/targets.json", "w", encoding="utf-8") as f:
        json.dump(targets_out, f, ensure_ascii=False, indent=2)

    with open(f"{OUTPUT_DIR}/associations.json", "w", encoding="utf-8") as f:
        json.dump(associations, f, ensure_ascii=False, indent=2)

    print(f"\nSaved to {OUTPUT_DIR}/:")
    print(f"  compounds.json    ({len(compounds_out)} records)")
    print(f"  targets.json     ({len(targets_out)} records)")
    print(f"  associations.json({len(associations)} records)")

    # Print unmatched summary
    print("\nTargets with 0 matches (need manual review):")
    for a in associations:
        if not a["matched_compound_ids"]:
            print(f"  [{a['target_id']}] {a['target_name']} → '{a['example_text']}'")


if __name__ == "__main__":
    main()
